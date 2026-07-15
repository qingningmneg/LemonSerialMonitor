from __future__ import annotations

import json
import shutil
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "artifacts" / "manual"
OUTPUT_PATH = OUTPUT_DIR / "Lemon串口监控-完整操作手册.docx"

PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
HEADER_FOOTER_DISTANCE = Inches(0.492)
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
EXPECTED_RENDERED_PAGE_COUNT = 15

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
CODE_BG = "F6F8FA"
CODE_BORDER = "D0D7DE"
NOTE_BG = "F4F6F9"
NOTE_BORDER = "2E74B5"
CAUTION_BG = "FFF4CE"
CAUTION_BORDER = "9A6700"
WARNING_BG = "FDECEC"
WARNING_BORDER = "C62828"
SUCCESS_BG = "EAF6EE"
SUCCESS_BORDER = "2E7D32"


def set_run_font(
    run,
    name: str = "Calibri",
    east_asia: str = "Microsoft YaHei UI",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, color: str = INK, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.right_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.header_distance = HEADER_FOOTER_DISTANCE
    section.footer_distance = HEADER_FOOTER_DISTANCE
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, 30, NAVY, True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 14, MUTED, False)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, BLUE, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13, BLUE, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 12, DARK_BLUE, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("Lemon串口监控  |  完整操作手册"), size=8.5, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    add_page_number(fp)

    doc.core_properties.title = "Lemon串口监控完整操作手册"
    doc.core_properties.subject = "安装、串口监控、复制、导出、AI 接入与完整卸载"
    doc.core_properties.author = "Lemon串口监控"
    doc.core_properties.last_modified_by = "Lemon串口监控"
    doc.core_properties.comments = "Lemon串口监控 0.1.0 完整操作手册"
    doc.core_properties.created = datetime(2026, 7, 15, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2026, 7, 15, tzinfo=timezone.utc)
    doc.core_properties.revision = 1
    doc.core_properties.keywords = "串口,监控,Windows,AI,MCP"


def normalize_extended_properties(path: Path, page_count: int) -> None:
    namespace = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    ET.register_namespace("", namespace)
    ET.register_namespace("vt", "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes")

    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "docProps/app.xml":
                root = ET.fromstring(data)
                values = {
                    "Application": "Lemon串口监控",
                    "AppVersion": "0.1.0",
                    "Company": "Lemon串口监控",
                    "Pages": str(page_count),
                }
                for name, value in values.items():
                    element = root.find(f"{{{namespace}}}{name}")
                    if element is None:
                        element = ET.SubElement(root, f"{{{namespace}}}{name}")
                    element.text = value
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(info, data)
    shutil.copyfile(temporary, path)
    temporary.unlink()


def append_abstract_numbering(doc: Document, abstract_id: int, bullet: bool) -> None:
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    rpr.append(fonts)
    level.append(rpr)
    abstract.append(level)
    first_num = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num, abstract)


def append_num_instance(doc: Document, abstract_id: int, num_id: int) -> None:
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    num.append(reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)


class Lists:
    def __init__(self, doc: Document) -> None:
        self.doc = doc
        self.next_num_id = 60
        append_abstract_numbering(doc, 31, bullet=True)
        append_abstract_numbering(doc, 32, bullet=False)
        append_num_instance(doc, 31, 41)

    def bullet(self, text: str):
        paragraph = self.doc.add_paragraph()
        apply_num(paragraph, 41)
        paragraph.add_run(text)
        return paragraph

    def checkbox(self, text: str):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
        paragraph.add_run(text)
        return paragraph

    def numbered(self, items: list[str]) -> None:
        num_id = self.next_num_id
        self.next_num_id += 1
        append_num_instance(self.doc, 32, num_id)
        for item in items:
            paragraph = self.doc.add_paragraph()
            apply_num(paragraph, num_id)
            paragraph.add_run(item)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shading = tcpr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tcpr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    trpr.append(marker)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must total {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for node in list(tblpr.findall(qn(tag))):
            tblpr.remove(node)
    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    tblpr.append(tblw)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    tblpr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.first_child_found_in("w:tcW")
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(header), size=9.5, bold=True, color=NAVY)
        shade_cell(cell, LIGHT_BLUE)
    set_repeat_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            set_run_font(paragraph.add_run(value), size=9.2, color=INK)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def set_paragraph_box(paragraph, fill: str, border: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    ppr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), border)
    borders.append(left)
    ppr.append(borders)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "160")
    indent.set(qn("w:right"), "100")
    ppr.append(indent)


def add_callout(doc: Document, label: str, text: str, kind: str = "note") -> None:
    palette = {
        "note": (NOTE_BG, NOTE_BORDER),
        "caution": (CAUTION_BG, CAUTION_BORDER),
        "warning": (WARNING_BG, WARNING_BORDER),
        "success": (SUCCESS_BG, SUCCESS_BORDER),
    }
    fill, border = palette[kind]
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_box(paragraph, fill, border)
    set_run_font(paragraph.add_run(f"{label}："), size=10.5, bold=True, color=border)
    set_run_font(paragraph.add_run(text), size=10.5, color=INK)


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    set_paragraph_box(paragraph, CODE_BG, CODE_BORDER)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        set_run_font(
            paragraph.add_run(line),
            name="Consolas",
            east_asia="Microsoft YaHei UI",
            size=8.5,
            color=INK,
        )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(paragraph.add_run(bold_prefix), bold=True, color=INK)
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(88)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("WINDOWS 串口被动监控  |  0.1.0"), size=10, bold=True, color=BLUE)
    doc.add_paragraph("Lemon串口监控", style="Title")
    doc.add_paragraph("完整操作手册", style="Subtitle")
    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_before = Pt(10)
    description.paragraph_format.space_after = Pt(72)
    set_run_font(
        description.add_run("安装 · 实时监控 · 复制 · 查找 · 导出 · AI/MCP · 完整卸载"),
        size=11,
        color=MUTED,
    )
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(4)
    set_run_font(metadata.add_run("兼容目标：Windows 10/11 x64，Windows Server 2019/2022/2025 x64"), size=9.5, color=MUTED)
    metadata2 = doc.add_paragraph()
    metadata2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(metadata2.add_run("文档版本：0.1.0  |  2026-07-15"), size=9.5, color=MUTED)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    lists = Lists(doc)
    add_cover(doc)
    add_page_break(doc)

    add_heading(doc, "使用导航", 1)
    add_table(
        doc,
        ["我要做什么", "直接阅读"],
        [
            ["第一次安装", "第 2–3 章：安装前准备、图形安装与重启"],
            ["马上开始监控", "第 1、4–5 章：快速开始、工具栏与三种视图"],
            ["复制和查找", "第 6–7 章：HEX/文本查找与八种复制格式"],
            ["交给 AI 读取", "第 10 章：MCP、CLI、分页和完整性"],
            ["排查读不到", "第 13 章：症状与处理顺序"],
            ["彻底卸载", "第 14 章：备份、完整卸载和重启续办"],
        ],
        [2500, 6860],
    )
    add_callout(
        doc,
        "最重要",
        "勾选端口后还必须点击“开始”，并让原业务软件在开始之后真正读写同一个 COM 端口。软件不会主动向设备发送测试数据。",
        "caution",
    )

    add_heading(doc, "1. 四步快速开始", 1)
    lists.numbered(
        [
            "打开原来使用串口的业务软件，让它按原方式连接设备。",
            "打开 Lemon串口监控，点击“刷新端口”，勾选目标 COM 端口。",
            "填写会话文件名并点击“开始”，再让原业务软件完成一次真实收发。",
            "看到事件后按需要查看、查找、复制；结束时先“停止”，再导出或关闭。",
        ]
    )
    add_callout(
        doc,
        "方向",
        "Write / TX 表示电脑发往设备；Read / RX 表示设备返回电脑；Ioctl 表示波特率、超时、线路控制等配置操作。",
    )

    add_heading(doc, "1.1 工作原理", 2)
    for item in (
        "桌面程序不打开被监控 COM 端口，原业务软件仍然是端口的实际使用者。",
        "Windows 内核过滤驱动复制已完成的串口 Read、Write 和受支持的控制事件。",
        "后台服务把事件写入受保护会话数据库，并提供桌面与 AI 命名管道。",
        "软件只读，不提供发送、注入、修改、阻断、重放或协议模拟。",
    ):
        lists.bullet(item)
    add_callout(
        doc,
        "边界",
        "不占用 COM 端口不等于系统中没有驱动参与。业务关键设备应先在可恢复环境验证，并保留完整卸载和重启窗口。",
        "warning",
    )

    add_heading(doc, "2. 安装前准备", 1)
    add_table(
        doc,
        ["检查项", "要求"],
        [
            ["兼容目标", "Windows 10/11 x64；Server 2019/2022/2025 x64"],
            ["权限", "本机管理员；安装和卸载会触发 UAC"],
            ["安全启动", "必须关闭；开启时安装程序会停止"],
            ["设备加密", "如启用 BitLocker，先确认恢复密钥可用"],
            ["维护窗口", "准备至少一次安装后重启和一次卸载后重启"],
            ["安装包", "只使用 GitHub Releases 正式文件并核对 SHA-256"],
        ],
        [1900, 7460],
    )
    add_heading(doc, "2.1 测试证书和 TESTSIGNING", 2)
    add_paragraph(
        doc,
        "0.1.0 驱动使用本地测试证书，不是微软正式发布签名。安装向导会展示完整协议，只有勾选接受后才能继续。安装包不包含证书私钥。",
    )
    for item in (
        "安装程序把随包公钥精确导入 LocalMachine Root 和 TrustedPublisher。",
        "把自签名测试证书加入系统证书库会改变本机信任边界；持有对应私钥的人签名的代码可能在本机被信任。安装包只携带公钥，不包含私钥。",
        "需要时自动启用 TESTSIGNING，并明确提示重启。",
        "不会关闭安全启动，不会修改 BitLocker，不会使用 nointegritychecks。",
        "卸载只移除本次安装拥有的证书；只有安装程序实际改变 TESTSIGNING 时才恢复关闭。",
        "安装文件没有微软信任链，也没有 RFC 3161 公共时间戳；首次运行仍可能显示 SmartScreen 或未知发布者。",
    ):
        lists.bullet(item)
    add_callout(
        doc,
        "先核对来源",
        "测试证书要等安装程序取得管理员权限后才能导入。只从本项目 GitHub Release 下载，并在运行前核对 SHA256SUMS.txt。",
        "warning",
    )

    add_heading(doc, "3. 图形化安装", 1)
    lists.numbered(
        [
            "双击“Lemon串口监控-安装程序-x64.exe”。",
            "遇到 SmartScreen 或 UAC 时核对下载来源和 SHA-256，再按 Windows 提示继续。",
            "阅读本地测试证书说明并勾选接受。",
            "选择桌面程序安装位置，默认是 C:\\Program Files\\Lemon串口监控。",
            "选择是否创建桌面快捷方式，在准备页核对安装模式和目标位置。",
            "点击“安装”；不要在运行过程中手工移动临时文件或结束 PowerShell 子进程。",
            "提示重启时先保存工作，再按提示重启；重启后再判断驱动是否可用。",
        ]
    )
    add_callout(
        doc,
        "一键安装做了什么",
        "安装程序自动校验每个载荷文件、签名和证书，安装驱动与后台服务，追加 Ports 类过滤器，写入受保护事务记录；中途失败会按已完成步骤回滚。",
        "success",
    )
    add_heading(doc, "3.1 安装位置", 2)
    add_table(
        doc,
        ["内容", "位置与说明"],
        [
            ["桌面程序", "向导中选择的位置"],
            ["后台核心", "%ProgramFiles% 下受保护内部目录，普通用户不可写"],
            ["会话与导出", "%ProgramData% 下受保护内部数据目录"],
            ["安装记录", "%ProgramData%\\LemonSerialMonitor\\Installer"],
            ["AI 租约", "授权用户 %LocalAppData%\\LemonSerialMonitor\\AI"],
        ],
        [2100, 7260],
    )
    add_heading(doc, "3.2 旧手工安装迁移", 2)
    add_paragraph(
        doc,
        "检测到早期同项目手工安装时，向导显示“迁移”。只有旧标记、备份、服务路径、驱动包、过滤器和证书全部相互验证后才会接管。迁移先创建受保护备份，中途失败会恢复旧文件和服务。身份不明确时会拒绝覆盖。",
    )

    add_heading(doc, "3.3 重启后验证", 2)
    lists.numbered(
        [
            "打开 Lemon串口监控并点击“刷新端口”。",
            "确认状态栏显示服务已连接，驱动不是 development fake source。",
            "勾选一个端口并点击“开始”。",
            "让原业务软件收发一次，确认列表出现 Read、Write 或 Ioctl。",
            "点击“停止”，完成一次 CSV 导出。",
        ]
    )
    add_callout(
        doc,
        "当前没有串口设备",
        "后台服务仍应保持运行。刷新后显示服务已连接、端口列表为空、驱动暂不可用属于正常状态；接入真实设备后再次刷新，不需要反复重装。",
        "note",
    )

    add_heading(doc, "4. 界面与捕获状态", 1)
    add_table(
        doc,
        ["控件", "作用与启用条件"],
        [
            ["刷新端口", "重新读取当前存在的串口，并刷新服务/驱动状态"],
            ["会话", "安全文件名；默认 capture.db，只影响下一次开始"],
            ["开始", "停止状态且至少勾选一个端口时启用"],
            ["暂停 / 继续", "暂停期间业务通信继续，但不生成监控副本"],
            ["停止", "结束捕获，不删除数据库和界面内容"],
            ["清空", "停止状态可用；确认后永久删除当前绑定会话"],
            ["导出", "停止状态可用；导出完整持久化会话"],
            ["复制数据", "选中事件后按指定格式写入剪贴板"],
            ["查找", "HEX 或文本向前/向后循环查找"],
        ],
        [2100, 7260],
    )
    add_callout(
        doc,
        "不要把清空当清屏",
        "清空会删除服务当前绑定数据库中的全部事件。只想开始新记录时，应停止后换一个会话文件名再开始。",
        "warning",
    )
    add_heading(doc, "4.1 完整捕获流程", 2)
    lists.numbered(
        [
            "让原业务软件保持原串口配置，不要为了监控改成另一个 COM 号。",
            "刷新端口并勾选目标端口；多端口可同时勾选。",
            "填写容易识别且不含目录的会话文件名。",
            "点击开始；状态改变后再操作硬件。",
            "需要短暂忽略一段通信时用暂停，恢复时用继续。",
            "完成后点击停止；按需要复制、查找、导出和记录完整性状态。",
        ]
    )

    add_heading(doc, "5. 三种数据视图", 1)
    add_heading(doc, "5.1 列表", 2)
    add_table(
        doc,
        ["列", "含义"],
        [
            ["序号 / 时间", "会话顺序与本地显示时间"],
            ["进程", "发起串口操作的进程名"],
            ["COM", "显示用端口名"],
            ["方向", "Read、Write、Ioctl 或 DropNotice"],
            ["操作 / 状态", "IOCTL 代码与 NTSTATUS"],
            ["长度", "完成长度；可能大于实际捕获长度"],
            ["标志", "输入/输出负载、截断、丢失等证据"],
            ["HEX / 文本", "同一负载的十六进制和 UTF-8 预览"],
        ],
        [1900, 7460],
    )
    add_paragraph(doc, "按住 Ctrl 或 Shift 可选择多行。选择列表事件后，Dump 和终端会联动到同一个底层事件。")

    add_heading(doc, "5.2 Dump", 2)
    add_paragraph(
        doc,
        "Dump 每行显示 16 字节：Offset 是十六进制偏移，HEX 是原始字节，ASCII 将不可打印字节显示为点。适合核对协议字段、长度、校验和与帧边界。",
    )

    add_heading(doc, "5.3 终端", 2)
    for item in (
        "只显示 Read/Write，忽略纯配置 Ioctl；只有 Ioctl 时终端为空是正常现象。",
        "Read 默认蓝色，Write 默认红色。",
        "支持 ANSI、UTF-7、UTF-8、UTF-16LE、UTF-16BE。",
        "可显示/隐藏时间、端口、方向，并切换自动换行和自动滚动。",
        "改变编码只影响后来到达的事件，不重新解释旧内容。",
    ):
        lists.bullet(item)

    add_heading(doc, "6. 查找", 1)
    add_heading(doc, "6.1 HEX 查找", 2)
    add_paragraph(doc, "每个字节写成两个十六进制字符并以空格分隔；?? 表示任意一个字节。")
    add_code(doc, "合法：01 03 00 FF\n合法：03 ?? FF\n非法：0x03\n非法：010300FF\n非法：GG")
    add_heading(doc, "6.2 文本查找", 2)
    add_paragraph(
        doc,
        "文本查找按 UTF-8 解释负载并忽略大小写。二进制协议、GBK 或其他编码优先使用 HEX 查找。上一个/下一个会从当前选择位置循环查找。",
    )

    add_heading(doc, "7. 复制数据", 1)
    add_table(
        doc,
        ["格式", "用途"],
        [
            ["HEX（空格）", "人工比对、协议记录，例如 01 03 00 FF"],
            ["HEX（紧凑）", "粘贴到要求连续十六进制的工具"],
            ["文本", "按 UTF-8 复制纯负载"],
            ["C 数组", "固件、C/C++ 单元测试与复现代码"],
            ["Python bytes", "Python 硬件测试与解析脚本"],
            ["TSV", "直接粘贴到 Excel/表格"],
            ["CSV", "保留事件字段，适合数据工具"],
            ["JSON", "保留事件字段，适合程序和 AI"],
        ],
        [2100, 7260],
    )
    lists.numbered(
        [
            "在列表选择一行或多行；Dump/终端选择的是对应完整事件。",
            "从“复制数据”右侧选择格式。",
            "点击“复制数据”或按 Ctrl+C。",
        ]
    )
    for item in (
        "Ctrl+C 使用当前格式。",
        "Ctrl+Shift+C 固定复制不带元数据的空格 HEX。",
        "HEX、文本、C 数组、Python bytes 多行时直接拼接负载。",
        "需要保留时间、端口、方向、进程和事件边界时用 TSV、CSV 或 JSON。",
    ):
        lists.bullet(item)

    add_heading(doc, "8. 会话与自动保存", 1)
    add_paragraph(
        doc,
        "点击开始后，事件持续写入受保护会话数据库，不需要再按一次保存。重复使用同名会话会继续写入；运行中修改文件名只影响下一次开始。",
    )
    add_callout(
        doc,
        "桌面与历史",
        "桌面界面以实时事件为主，不把历史数据库重新加载回三个视图。历史分页读取、等待和 JSON/JSONL 导出由 AI/命令行接口提供。",
        "note",
    )
    add_heading(doc, "8.1 导出", 2)
    lists.numbered(
        [
            "点击停止。",
            "填写导出文件名，选择 CSV、TXT 或 RAW。",
            "点击导出，成功后在状态栏查看完整输出位置。",
            "将要长期保留的文件复制到本软件管理目录之外。",
        ]
    )
    add_table(
        doc,
        ["格式", "内容与适用场景"],
        [
            ["CSV", "结构化事件字段，适合 Excel、数据库和脚本"],
            ["TXT", "便于人工查看的文本表示"],
            ["RAW", "按顺序拼接原始负载，不保留事件边界和元数据"],
        ],
        [1800, 7560],
    )
    add_callout(
        doc,
        "导出范围",
        "导出的是当前服务绑定会话的全部持久化记录，不是当前选中的几行，也不一定等于界面仍保留的可见行。",
        "caution",
    )

    add_heading(doc, "9. 状态、容量和证据", 1)
    add_table(
        doc,
        ["层次", "边界"],
        [
            ["单事件", "最多捕获 4096 字节；更长事件带截断标志"],
            ["界面待处理", "最多 10,000 条；溢出计入丢失"],
            ["列表", "最近 100,000 行；旧行仍可在持久化会话中"],
            ["终端", "约 2 MiB 可见文本；超出后裁掉最旧片段"],
            ["AI 页面", "1–1000 条；通过游标和回执续读"],
        ],
        [1900, 7460],
    )
    add_paragraph(
        doc,
        "状态栏的“丢失 = 0”不能单独证明全链路无丢包。严谨分析还要检查截断、驱动丢弃、服务丢弃、序号缺口、AI 完整性字段、原发送端/接收端日志和文件哈希。",
    )

    add_heading(doc, "10. AI、MCP 与命令行", 1)
    add_paragraph(
        doc,
        "AI 接口使用本机命名管道，不监听 HTTP/TCP。服务校验 Windows 用户、登录会话、客户端绝对路径和 SHA-256。接口不直接打开 COM 端口。",
    )
    add_heading(doc, "10.1 MCP 配置", 2)
    add_code(
        doc,
        '{\n  "mcpServers": {\n    "lemon-serial-monitor": {\n      "command": "C:\\\\Program Files\\\\Lemon串口监控\\\\ai\\\\Lemon.SerialMonitor.AI.exe",\n      "args": ["mcp"]\n    }\n  }\n}',
    )
    add_paragraph(doc, "安装到其他位置时必须改成实际绝对路径。连接成功应列出 11 个工具和 4 个资源。")
    add_paragraph(
        doc,
        "没有接入串口设备时，状态工具仍应返回服务可用，driverState 可为不可用，端口工具返回空数组。这是正常的无设备状态，AI 不应把它解释成后台服务崩溃。",
    )
    add_heading(doc, "10.2 推荐调用顺序", 2)
    lists.numbered(
        [
            "lemon_get_status：确认服务、驱动、捕获和完整性状态。",
            "lemon_list_ports：取得 16 位十六进制 deviceId。",
            "lemon_start_capture：开始并保存 leaseId。",
            "让原业务软件操作硬件。",
            "lemon_list_sessions：取得安全 sessionId。",
            "lemon_read_events / lemon_wait_events：按游标读取。",
            "每页检查 integrity、warnings、nextCursor 和 resumeReceipt。",
            "lemon_stop_capture：使用原 leaseId 停止。",
        ]
    )
    add_heading(doc, "10.3 CLI 快速命令", 2)
    add_code(
        doc,
        "$lemon = 'C:\\Program Files\\Lemon串口监控\\ai\\Lemon.SerialMonitor.AI.exe'\n"
        "& $lemon status --json\n"
        "& $lemon ports --json\n"
        "& $lemon sessions list --limit 100 --json\n"
        "& $lemon schema --json",
    )
    add_code(
        doc,
        "& $lemon capture start --device-id 0000000000000011 --label board-test --json\n"
        "& $lemon capture pause --lease-id '<leaseId>' --json\n"
        "& $lemon capture resume --lease-id '<leaseId>' --json\n"
        "& $lemon events read --session-id '<sessionId>' --limit 100 --include-hex --json\n"
        "& $lemon events wait --session-id '<sessionId>' --cursor '<nextCursor>' --resume-receipt '<resumeReceipt>' --limit 100 --timeout-seconds 30 --include-hex --jsonl\n"
        "& $lemon export --session-id '<sessionId>' --format jsonl --label board-test --json\n"
        "& $lemon capture stop --lease-id '<leaseId>' --json",
    )

    add_heading(doc, "10.4 CLI 退出码", 2)
    add_table(
        doc,
        ["退出码", "含义"],
        [
            ["0", "成功"],
            ["2", "参数或输入错误"],
            ["3", "访问被拒绝或协议不兼容"],
            ["4", "后台服务、驱动或设备不可用"],
            ["5", "捕获冲突，或租约无效/过期"],
            ["6", "数据缺口、完整性未知或连续性未证明"],
            ["7", "超时或操作取消"],
            ["10", "未预期运行错误"],
        ],
        [1700, 7660],
    )

    add_heading(doc, "10.5 完整性判定", 2)
    add_table(
        doc,
        ["字段", "判定"],
        [
            ["completeForReturnedRange", "只有 true 才能声明返回范围完整"],
            ["continuityProven", "游标续读连续性是否有证据"],
            ["driverDropped / serviceDropped", "驱动或服务是否丢弃事件"],
            ["truncationSeen", "返回范围是否包含截断事件"],
            ["gapDetected", "是否发现序号或提交缺口"],
            ["warnings", "必须保留并向使用者说明的风险"],
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "AI 不得猜数据",
        "完整性不是 true 时，AI 必须明确说明证据不足，不能把缺失字节、设备响应或时序补写成事实。",
        "warning",
    )

    add_heading(doc, "11. Windows Server", 1)
    add_table(
        doc,
        ["系统", "安装内容"],
        [
            ["Server 2019/2022/2025 桌面体验", "桌面程序、服务、驱动、AI、文档和快捷方式"],
            ["Server 2019/2022/2025 Server Core", "服务、驱动、AI/CLI、文档；不安装 WPF"],
        ],
        [3300, 6060],
    )
    add_paragraph(
        doc,
        "Server Core 使用 MCP 或 JSON CLI 操作。安装器只接受已知正式构建，未知 Server 版本会停止。测试签名、安全启动、重启和完整卸载要求与桌面系统相同。",
    )
    add_callout(
        doc,
        "验证范围",
        "Windows Server 2022/2025 只完成 GitHub 托管桌面 runner 的平台、托管测试和安装契约检查，未装载内核驱动；Server Core 只有布局契约测试，Server 2019 自托管任务未执行。0.1.0 没有任何 Server 驱动端到端实测，重要环境必须先在同版本测试机完成全流程验证。",
        "caution",
    )

    add_heading(doc, "12. 数据备份与更新", 1)
    add_heading(doc, "12.1 导出备份", 2)
    for item in (
        "停止捕获后导出 CSV/TXT/RAW；AI 还可导出 JSON/JSONL。",
        "把要保留的文件复制到本软件管理目录之外。",
        "完整数据库备份必须包括同名 DB、-wal 和 -shm；不要只复制主文件后删除 WAL。",
        "敏感串口负载可能包含密钥、令牌、客户数据或固件，公开前先审查。",
    ):
        lists.bullet(item)
    add_heading(doc, "12.2 更新", 2)
    lists.numbered(
        [
            "停止监控并备份需要保留的数据。",
            "使用旧版本卸载程序完整卸载。",
            "按提示重启并等待清理完成。",
            "安装新版本并再次重启/验证。",
        ]
    )
    add_paragraph(doc, "0.1.0 不支持在已有新式安装上原地覆盖。不要直接覆盖正在加载的 SYS、后台服务 EXE 或受保护安装记录。")

    add_heading(doc, "13. 常见故障", 1)
    add_table(
        doc,
        ["症状", "先做什么", "仍失败时"],
        [
            ["完全没数据", "刷新、勾选、开始；让原业务软件真实通信", "核对 COM 号和状态栏错误"],
            ["端口为空", "确认设备管理器是否存在可用串口", "无设备时服务运行、驱动暂不可用属于正常状态"],
            ["终端空", "看列表是否只有 Ioctl", "确认有 Read/Write 正文"],
            ["服务未连接", "安装后重启；检查服务状态", "查看 Windows 应用/系统事件"],
            ["驱动未就绪", "检查重启、TESTSIGNING、安全启动", "检查证书、策略和驱动事件"],
            ["终端乱码", "按协议切换编码或改看 HEX", "保存原始字节而非文本猜测"],
            ["导出不可用", "先停止并填写文件名", "检查状态栏错误"],
            ["AI 拒绝", "使用安装目录原始 EXE 和授权用户", "核对绝对路径、登录会话和服务"],
        ],
        [1800, 3300, 4260],
    )
    add_heading(doc, "13.1 管理员状态命令", 2)
    add_code(
        doc,
        "sc.exe query CommMonitorService\n"
        "sc.exe qc CommMonitorService\n"
        "sc.exe query CommMonitorFilter\n"
        "bcdedit.exe /enum '{current}' | Select-String testsigning",
    )
    add_callout(
        doc,
        "原业务串口异常",
        "优先停止监控、关闭客户端、记录时间和设备信息，然后使用完整卸载并重启。不要反复手改 UpperFilters、删除服务或覆盖驱动。",
        "warning",
    )

    add_heading(doc, "14. 完整卸载", 1)
    add_callout(
        doc,
        "不可恢复",
        "完整卸载会永久删除本软件产生的全部会话、导出、设置、日志、缓存和 AI 状态。需要保留的数据必须先复制到软件目录之外。",
        "warning",
    )
    lists.numbered(
        [
            "打开 Windows 设置 → 应用 → 已安装的应用。",
            "找到 Lemon串口监控并点击卸载。",
            "阅读数据删除警告并确认，允许管理员权限。",
            "卸载程序会先关闭本软件桌面程序、AI 客户端和后台服务，再清理驱动、过滤器、证书、文件和数据。",
            "提示重启时先保存工作，再按提示重启；卸载程序会自动继续并核验残留。",
        ]
    )
    add_paragraph(
        doc,
        "正常用户态文件不应只因本软件自身仍在运行而要求重启；内核驱动、Ports 类设备栈、启动策略或 Windows 文件锁仍可能必须在重启后完成清理。",
    )
    add_heading(doc, "14.1 卸载会清理什么", 2)
    for item in (
        "桌面程序、AI 客户端、文档、开始菜单和桌面快捷方式。",
        "后台服务、驱动服务、精确 Driver Store 包和 Ports 类过滤器条目。",
        "本次安装实际拥有的测试证书。",
        "会话、导出、设置、日志、缓存、AI 状态和迁移备份。",
        "安装记录、卸载续办任务和安装程序自身文件。",
    ):
        lists.bullet(item)
    add_paragraph(
        doc,
        "卸载只按受保护记录和精确对象身份执行。服务路径、驱动 INF、证书指纹或文件身份不一致时会停止对应删除，避免误删其他软件。",
    )
    add_heading(doc, "14.2 TESTSIGNING 恢复", 2)
    add_paragraph(
        doc,
        "如果安装程序是本次启用 TESTSIGNING 的，卸载会恢复关闭并要求重启；如果安装前已经开启，卸载不改变原有策略。",
    )

    add_heading(doc, "15. AI 工具与资源索引", 1)
    add_table(
        doc,
        ["类型", "名称"],
        [
            ["状态/端口", "lemon_get_status, lemon_list_ports"],
            ["捕获", "lemon_start_capture, lemon_pause_capture, lemon_resume_capture, lemon_stop_capture"],
            ["读取", "lemon_list_sessions, lemon_read_events, lemon_wait_events"],
            ["导出/描述", "lemon_export_session, lemon_get_schema"],
            ["资源", "lemon://docs/ai-interface"],
            ["资源", "lemon://schema/capture-event, errors, integrity"],
        ],
        [2100, 7260],
    )

    add_page_break(doc)
    add_heading(doc, "16. 最终检查清单", 1)
    for item in (
        "□ 安装包来自正式发布页，SHA-256 一致。",
        "□ 已理解测试证书、TESTSIGNING、安全启动和重启影响。",
        "□ 监控前刷新、勾选并点击开始。",
        "□ 无设备时已确认服务保持运行，接入设备后再刷新。",
        "□ 原业务软件在开始后发生真实读写。",
        "□ 结束时先停止，再复制或导出。",
        "□ 严谨分析检查截断、丢弃、缺口和 AI 完整性。",
        "□ 卸载前把需要保留的数据移到软件目录之外。",
        "□ 卸载后按提示重启并等待续办完成。",
    ):
        lists.checkbox(item)
    add_callout(
        doc,
        "证据原则",
        "任何监控界面都不能替代原发送端和接收端日志。做故障定责、协议验证或无损验收时，应同时保留双方原始日志、会话/导出、状态输出和文件哈希。",
        "note",
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    normalize_extended_properties(OUTPUT_PATH, EXPECTED_RENDERED_PAGE_COUNT)
    return doc


def audit_document(path: Path) -> dict[str, object]:
    if not path.is_file() or path.stat().st_size < 30_000:
        raise RuntimeError(f"DOCX missing or unexpectedly small: {path}")

    reopened = Document(path)
    section = reopened.sections[0]
    expected = {
        "page_width": int(PAGE_WIDTH),
        "page_height": int(PAGE_HEIGHT),
        "top_margin": int(MARGIN),
        "right_margin": int(MARGIN),
        "bottom_margin": int(MARGIN),
        "left_margin": int(MARGIN),
    }
    actual = {name: int(getattr(section, name)) for name in expected}
    if actual != expected:
        raise RuntimeError(f"page geometry mismatch: {actual}")

    for table_index, table in enumerate(reopened.tables, start=1):
        tblpr = table._tbl.tblPr
        tblw = tblpr.find(qn("w:tblW"))
        tblind = tblpr.find(qn("w:tblInd"))
        layout = tblpr.find(qn("w:tblLayout"))
        if tblw is None or tblw.get(qn("w:w")) != str(TABLE_WIDTH_DXA):
            raise RuntimeError(f"table {table_index} width mismatch")
        if tblind is None or tblind.get(qn("w:w")) != str(TABLE_INDENT_DXA):
            raise RuntimeError(f"table {table_index} indent mismatch")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise RuntimeError(f"table {table_index} is not fixed layout")

    title = next((p for p in reopened.paragraphs if p.style.name == "Title"), None)
    if title is None or title.text != "Lemon串口监控":
        raise RuntimeError("title is missing or incorrect")
    if title._p.get_or_add_pPr().find(qn("w:pBdr")) is not None:
        raise RuntimeError("title paragraph unexpectedly contains a border")

    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        extended_xml = archive.read("docProps/app.xml").decode("utf-8")
        package_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        )
    required = (
        "Lemon串口监控",
        "八种复制格式",
        "MCP",
        "完整卸载",
        "Windows Server",
        "completeForReturnedRange",
        "RFC 3161",
        "后台服务仍应保持运行",
        "不支持在已有新式安装上原地覆盖",
        "Server 2019 自托管任务未执行",
        "卸载程序会先关闭本软件桌面程序",
    )
    for phrase in required:
        if phrase not in package_xml:
            raise RuntimeError(f"required manual phrase missing: {phrase}")
    forbidden = "CommMonitor " + "串口监控精灵"
    if forbidden in package_xml:
        raise RuntimeError("retired public product name is present")
    extended_root = ET.fromstring(extended_xml)
    extended_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    pages = extended_root.find(f"{{{extended_namespace}}}Pages")
    application = extended_root.find(f"{{{extended_namespace}}}Application")
    if pages is None or pages.text != str(EXPECTED_RENDERED_PAGE_COUNT):
        raise RuntimeError("extended page count does not match the verified render")
    if application is None or application.text != "Lemon串口监控":
        raise RuntimeError("extended application metadata is incorrect")
    if "w:numFmt w:val=\"bullet\"" not in numbering_xml:
        raise RuntimeError("real bullet numbering definition is missing")
    for paragraph in reopened.paragraphs:
        if paragraph.text.startswith("□"):
            numpr = paragraph._p.get_or_add_pPr().find(qn("w:numPr"))
            if numpr is not None:
                raise RuntimeError("checklist item unexpectedly contains a bullet number")

    headings = sum(p.style.name.startswith("Heading") for p in reopened.paragraphs)
    return {
        "path": str(path),
        "bytes": path.stat().st_size,
        "paragraphs": len(reopened.paragraphs),
        "tables": len(reopened.tables),
        "headings": headings,
        "page_geometry": actual,
        "preset": "compact_reference_guide",
        "header_template": "editorial_cover",
        "structural_audit": "PASS",
    }


if __name__ == "__main__":
    build_document()
    print(json.dumps(audit_document(OUTPUT_PATH), ensure_ascii=False, indent=2))
